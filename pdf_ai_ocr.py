"""
PDF AI识别工具 - 仅支持LM Studio
开源项目: https://github.com/bbblq/pdf-ai-ocr
"""

import os
import re
import base64
import requests
import fitz
from pathlib import Path
from flask import Flask, request, jsonify, send_file
from flask import render_template_string, Response
import uuid
import json
import time
import threading
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 500 * 1024 * 1024  # 500MB max
app.config["UPLOAD_FOLDER"] = "uploads"
app.config["OUTPUT_FOLDER"] = "outputs"

os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
os.makedirs(app.config["OUTPUT_FOLDER"], exist_ok=True)

tasks = {}

# LM Studio 配置
LM_STUDIO_URL = "http://localhost:1234"

# ============== 预设提示词 ==============
PRESET_PROMPTS = {
    "contract": {
        "name": "📜 合同文档",
        "prompt": "【严格OCR任务】你只能输出图片中的原始文字，不要做任何解释、总结、改写或发挥。直接逐字输出看到的所有文字内容，包括合同条款、金额、日期、签名等。如有表格，用制表符分隔。禁止添加任何说明。",
    },
    "invoice": {
        "name": "🧾 发票收据",
        "prompt": "【严格OCR任务】你只能输出图片中的原始文字，不要做任何解释、总结、改写或发挥。直接逐字输出看到的所有数字和文字。如有表格，用制表符分隔。禁止添加任何说明。",
    },
    "idcard": {
        "name": "🪪 身份证/证件",
        "prompt": "【严格OCR任务】你只能输出图片中的原始文字，不要做任何解释、总结、改写或发挥。直接逐字输出看到的所有字符。禁止添加任何说明。",
    },
    "book": {
        "name": "📚 书籍资料",
        "prompt": "【严格OCR任务】你只能输出图片中的原始文字，不要做任何解释、总结、改写或发挥。保持原有排版格式，逐字输出。禁止添加任何说明。",
    },
    "general": {
        "name": "📝 通用文档",
        "prompt": "【严格OCR任务】你只能输出图片中的原始文字，不要做任何解释、总结、改写或发挥。直接逐字输出看到的所有内容。禁止添加任何说明。",
    },
    "table": {
        "name": "📊 表格数据",
        "prompt": "【严格OCR任务】你只能输出图片中的原始表格数据，不要做任何解释、总结、改写或发挥。用制表符分隔各列，保持行列结构。禁止添加任何说明。",
    },
}


# ============== API接口函数 ==============


def get_models(url):
    """获取LM Studio模型列表"""
    try:
        response = requests.get(f"{url}/v1/models", timeout=5)
        if response.status_code == 200:
            models = response.json().get("data", [])
            return [m["id"] for m in models]
    except:
        pass
    return []


def call_vl_model(url, model_name, image_base64, prompt, max_retries=3):
    """调用LM Studio的VL模型"""
    for attempt in range(max_retries):
        try:
            response = requests.post(
                f"{url}/v1/chat/completions",
                json={
                    "model": model_name,
                    "messages": [
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/png;base64,{image_base64}"
                                    },
                                },
                                {"type": "text", "text": prompt},
                            ],
                        }
                    ],
                    "max_tokens": 4000,
                    "temperature": 0.1,
                },
                timeout=180,
            )

            if response.status_code == 200:
                result = response.json()
                msg = result["choices"][0]["message"]
                content = msg.get("content") or msg.get("reasoning_content") or ""
                if not content:
                    return "[错误: 模型未返回有效内容]"
                return content
            elif (
                "failed to process image" in response.text
                or "does not support image" in response.text
            ):
                return f"[错误: 图片处理失败]"
            else:
                print(
                    f"  LM Studio API错误: {response.status_code} - {response.text[:200]}"
                )

        except requests.exceptions.Timeout:
            print(f"  请求超时")
        except Exception as e:
            print(f"  错误: {e}")

        if attempt < max_retries - 1:
            time.sleep(3)

    return "[错误: 无法从LM Studio获取响应]"


# ============== Flask路由 ==============


@app.route("/")
def index():
    return render_template_string(HTML_TEMPLATE)


@app.route("/api/models", methods=["GET"])
def get_models_api():
    """获取模型列表"""
    url = request.args.get("url", LM_STUDIO_URL)
    models = get_models(url)

    # 过滤视觉模型
    vision_keywords = [
        "vl",
        "vision",
        "llava",
        "qwen2-vl",
        "qwen3-vl",
        "llama3.2-vision",
    ]
    vision_models = [m for m in models if any(k in m.lower() for k in vision_keywords)]

    return jsonify(
        {
            "success": True if models else False,
            "url": url,
            "all_models": models,
            "vision_models": vision_models if vision_models else models[:10],
        }
    )


@app.route("/api/presets", methods=["GET"])
def get_preset_prompts():
    """获取预设提示词列表"""
    return jsonify(
        {
            "success": True,
            "presets": [
                {"id": k, "name": v["name"], "prompt": v["prompt"]}
                for k, v in PRESET_PROMPTS.items()
            ],
        }
    )


@app.route("/api/upload", methods=["POST"])
def upload_file():
    """上传PDF文件"""
    if "file" not in request.files:
        return jsonify({"success": False, "error": "没有文件"})

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"success": False, "error": "文件名为空"})

    if not file.filename.lower().endswith(".pdf"):
        return jsonify({"success": False, "error": "只支持PDF文件"})

    task_id = str(uuid.uuid4())[:8]
    filename = f"{task_id}_{file.filename}"
    filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    file.save(filepath)

    try:
        doc = fitz.open(filepath)
        total_pages = len(doc)
        doc.close()
    except Exception as e:
        os.remove(filepath)
        return jsonify({"success": False, "error": f"PDF读取失败: {str(e)}"})

    tasks[task_id] = {
        "id": task_id,
        "filename": file.filename,
        "filepath": filepath,
        "total_pages": total_pages,
        "current_page": 0,
        "status": "ready",
        "results": {},
        "errors": [],
        "output_file": None,
        "model": None,
        "stop_flag": False,
    }

    return jsonify(
        {
            "success": True,
            "task_id": task_id,
            "filename": file.filename,
            "total_pages": total_pages,
        }
    )


@app.route("/api/start", methods=["POST"])
def start_processing():
    """开始处理PDF"""
    data = request.json
    task_id = data.get("task_id")
    model_name = data.get("model")
    url = data.get("url", LM_STUDIO_URL)
    prompt = data.get("prompt", "")

    if task_id not in tasks:
        return jsonify({"success": False, "error": "任务不存在"})

    if not model_name:
        return jsonify({"success": False, "error": "请选择模型"})

    task = tasks[task_id]
    task["model"] = model_name

    if task["status"] == "running":
        return jsonify({"success": False, "error": "任务已在运行中"})

    thread = threading.Thread(
        target=process_pdf, args=(task_id, url, model_name, prompt)
    )
    thread.daemon = True
    thread.start()

    return jsonify({"success": True, "message": "处理已开始"})


@app.route("/api/stop/<task_id>", methods=["POST"])
def stop_processing(task_id):
    """停止处理"""
    if task_id not in tasks:
        return jsonify({"success": False, "error": "任务不存在"})

    task = tasks[task_id]
    if task["status"] == "running":
        task["stop_flag"] = True
        return jsonify({"success": True, "message": "已发送停止信号"})

    return jsonify({"success": False, "error": "任务不在运行中"})


@app.route("/api/status/<task_id>")
def get_status(task_id):
    if task_id not in tasks:
        return jsonify({"success": False, "error": "任务不存在"})

    task = tasks[task_id]
    return jsonify(
        {
            "success": True,
            "status": task["status"],
            "current_page": task["current_page"],
            "total_pages": task["total_pages"],
            "progress": round(task["current_page"] / task["total_pages"] * 100, 1)
            if task["total_pages"] > 0
            else 0,
            "errors": task["errors"][-5:],
        }
    )


@app.route("/api/stream/<task_id>")
def stream_status(task_id):
    def generate():
        last_page = -1
        while True:
            if task_id not in tasks:
                yield f"data: {json.dumps({'done': True, 'error': '任务不存在'})}\n\n"
                break

            task = tasks[task_id]

            current = task["current_page"] or 0
            total = task["total_pages"] or 1

            if current != last_page:
                last_page = current
                progress = round(current / total * 100, 1) if total > 0 else 0
                yield f"data: {json.dumps({'current_page': current, 'total_pages': total, 'progress': progress, 'status': task['status']})}\n\n"

            if task["status"] in ["completed", "error"]:
                # 最终状态也要发送页面信息
                yield f"data: {json.dumps({'done': True, 'current_page': current, 'total_pages': total, 'progress': 100 if task['status'] == 'completed' else 0, 'status': task['status']})}\n\n"
                break

            time.sleep(0.5)

    return Response(generate(), mimetype="text/event-stream")


@app.route("/api/download/<task_id>")
def download_result(task_id):
    """下载结果文件，支持docx和md格式"""
    fmt = request.args.get("format", "docx")

    if task_id not in tasks:
        return jsonify({"success": False, "error": "任务不存在"})

    task = tasks[task_id]

    if fmt == "md":
        # 生成Markdown文件
        output_file = save_results_to_markdown(task)
    else:
        # 生成Word文件
        output_file = save_results_to_word(task)

    if not output_file or not os.path.exists(output_file):
        return jsonify({"success": False, "error": "文件不存在"})

    filename = f"{task['id']}_output.{fmt}"
    return send_file(output_file, as_attachment=True, download_name=filename)


@app.route("/api/preview/<task_id>/<int:page>")
def preview_page(task_id, page):
    if task_id not in tasks:
        return jsonify({"success": False, "error": "任务不存在"})

    task = tasks[task_id]
    page_result = task["results"].get(page)

    if page_result is None:
        return jsonify({"success": False, "error": "该页未处理或无结果"})

    return jsonify({"success": True, "page": page, "content": page_result})


# ============== 处理函数 ==============


def extract_page_as_image(pdf_path, page_num, scale=1.0):
    doc = fitz.open(pdf_path)
    page = doc[page_num]
    mat = fitz.Matrix(scale, scale)
    pix = page.get_pixmap(matrix=mat)
    img_bytes = pix.tobytes("png")
    doc.close()
    return img_bytes


def is_markdown_table(lines, start_idx):
    """检测是否是markdown表格（|分隔）"""
    if start_idx >= len(lines):
        return False
    line = lines[start_idx].strip()
    # 检查是否是表格行（|分隔）
    if "|" in line:
        # 下一行应该是 --- 分隔行
        if start_idx + 1 < len(lines):
            next_line = lines[start_idx + 1].strip()
            if next_line.startswith("|") and "---" in next_line:
                return True
    return False


def parse_markdown_table(lines, start_idx):
    """解析markdown表格，返回行列表和结束索引"""
    table_rows = []
    i = start_idx
    # 跳过表头后的 --- 行
    if i + 1 < len(lines) and "---" in lines[i + 1]:
        i += 1

    while i < len(lines):
        line = lines[i].strip()
        if not line or "|" not in line:
            break
        # 解析表格行
        cols = [c.strip() for c in line.split("|") if c.strip()]
        if cols:
            table_rows.append(cols)
        i += 1

    return table_rows, i - 1


def clean_text(text):
    """清理文本，移除多余空格、HTML标签和处理中文间隔问题"""
    # 移除HTML标签
    text = re.sub(r"<[^>]+>", "", text)
    # 处理中文间隔问题：如果一个中文字和下一个中文字之间有空格，去掉
    # 匹配中文和其他CJK字符，移除它们之间的空格
    text = re.sub(r"([\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])", r"\1", text)
    text = re.sub(r"(?<=[\u4e00-\u9fff])\s+([\u4e00-\u9fff])", r"\1", text)
    # 移除连续空格
    text = re.sub(r" +", " ", text)
    # 移除换行前后的空格
    text = re.sub(r"\n\s+", "\n", text)
    text = re.sub(r"\s+\n", "\n", text)
    return text.strip()


def add_formatted_paragraph(doc, text, font_size=11):
    """添加段落，统一字体"""
    text = clean_text(text)
    if not text:
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "黑体")
    return p


def add_formatted_table(doc, rows, font_size=11):
    """添加Word表格，统一字体"""
    if not rows:
        return

    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= max_cols:
                break
            cell = table.cell(i, j)
            cell.text = clean_text(cell_text)
            # 设置单元格字体
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "黑体")


def save_results_to_word(task):
    """保存结果到Word文档"""
    output_file = os.path.join(
        app.config["OUTPUT_FOLDER"],
        f"{task['id']}_output.docx",
    )

    doc = Document()

    # 标题
    title = doc.add_heading(f"PDF识别结果 - {task['filename']}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "黑体")

    # 文档信息
    info_para = doc.add_paragraph()
    run1 = info_para.add_run(f"总页数: {task['total_pages']}")
    run1.bold = True
    run1.font.size = Pt(11)
    run1.font.name = "Times New Roman"
    run1._element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "黑体")
    run2 = info_para.add_run(f"\n模型: {task.get('model', 'N/A')}")
    run2.font.size = Pt(11)
    run2.font.name = "Times New Roman"
    run2._element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "黑体")

    doc.add_paragraph("─" * 50)

    # 按页添加内容
    for page_num in sorted(task["results"].keys()):
        # 页码标题
        page_heading = doc.add_heading(f"第 {page_num + 1} 页", level=1)
        page_heading.runs[0].font.size = Pt(12)
        page_heading.runs[0].font.name = "Times New Roman"
        page_heading.runs[0]._element.rPr.rFonts.set(
            docx.oxml.ns.qn("w:eastAsia"), "黑体"
        )

        # 内容
        content = task["results"][page_num]
        if "[错误:" in content or "[Error:" in content or "识别失败" in content:
            add_formatted_paragraph(doc, content)
        else:
            # 解析内容，处理markdown表格
            lines = content.split("\n")
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                if not line:
                    i += 1
                    continue

                # 检查是否是表格开始
                if is_markdown_table(lines, i):
                    # 解析markdown表格
                    table_rows, end_idx = parse_markdown_table(lines, i)
                    if table_rows:
                        add_formatted_table(doc, table_rows)
                    i = end_idx + 1
                else:
                    # 非表格行
                    add_formatted_paragraph(doc, line)
                    i += 1

        doc.add_paragraph()

    try:
        doc.save(output_file)
        print(f"文件已保存: {output_file}")
        return output_file
    except Exception as e:
        print(f"保存文件失败: {e}")
        raise e


def save_results_to_markdown(task):
    """保存结果到Markdown文件"""
    output_file = os.path.join(
        app.config["OUTPUT_FOLDER"],
        f"{task['id']}_output.md",
    )

    with open(output_file, "w", encoding="utf-8") as f:
        f.write(f"# {task['filename']}\n\n")
        f.write(f"**总页数**: {task['total_pages']}\n")
        f.write(f"**模型**: {task.get('model', 'N/A')}\n\n")
        f.write("---\n\n")

        for page_num in sorted(task["results"].keys()):
            f.write(f"## 第 {page_num + 1} 页\n\n")
            f.write(clean_text(task["results"][page_num]))
            f.write("\n\n---\n\n")

    print(f"文件已保存: {output_file}")
    return output_file


def process_pdf(task_id, url, model_name, prompt):
    task = tasks[task_id]
    task["status"] = "running"
    task["current_page"] = 0
    task["stop_flag"] = False

    try:
        # 先测试模型是否支持图片
        print(f"  测试模型 {model_name} 是否支持图片...")
        test_img = extract_page_as_image(task["filepath"], 0, scale=0.5)
        test_base64 = base64.b64encode(test_img).decode("utf-8")
        test_result = call_vl_model(url, model_name, test_base64, "hi", max_retries=1)

        if (
            "不支持" in test_result
            or "does not support" in test_result
            or "Cannot read" in test_result
        ):
            task["status"] = "error"
            task["errors"].append(f"模型 {model_name} 不支持图片输入，请使用VL模型")
            return

        print(f"  模型测试通过，开始处理...")

        for i in range(task["total_pages"]):
            # 检查停止标志
            if task.get("stop_flag"):
                print(f"  用户停止处理，第{i + 1}页停止")
                task["status"] = "stopped"
                return

            task["current_page"] = i + 1  # 1-indexed for display

            try:
                img_bytes = extract_page_as_image(task["filepath"], i, scale=1.0)
                img_base64 = base64.b64encode(img_bytes).decode("utf-8")
                result = call_vl_model(url, model_name, img_base64, prompt)

                if (
                    "图片处理失败" in result
                    or "不支持" in result
                    or "Cannot read" in result
                ):
                    print(f"  第{i + 1}页处理失败，尝试更低清晰度...")
                    if task.get("stop_flag"):
                        task["status"] = "stopped"
                        return
                    img_bytes = extract_page_as_image(task["filepath"], i, scale=0.75)
                    img_base64 = base64.b64encode(img_bytes).decode("utf-8")
                    result = call_vl_model(url, model_name, img_base64, prompt)

                if (
                    "图片处理失败" in result
                    or "不支持" in result
                    or "Cannot read" in result
                ):
                    print(f"  第{i + 1}页再次失败，尝试最低清晰度...")
                    if task.get("stop_flag"):
                        task["status"] = "stopped"
                        return
                    img_bytes = extract_page_as_image(task["filepath"], i, scale=0.5)
                    img_base64 = base64.b64encode(img_bytes).decode("utf-8")
                    result = call_vl_model(url, model_name, img_base64, prompt)

                # 检查是否是错误结果
                error_patterns = [
                    "[错误:",
                    "[Error:",
                    "不支持",
                    "does not support",
                    "Cannot read",
                    "图片处理失败",
                ]
                is_error = any(pattern in result for pattern in error_patterns)

                if is_error:
                    task["errors"].append(f"第{i + 1}页: {result[:100]}")
                    task["results"][i] = f"[第{i + 1}页识别失败]"
                else:
                    task["results"][i] = result

            except Exception as e:
                task["errors"].append(f"第{i + 1}页: {str(e)}")
                task["results"][i] = f"[第{i + 1}页识别失败]"

            time.sleep(1.0)

        task["output_file"] = save_results_to_word(task)
        task["status"] = "completed"

    except Exception as e:
        task["status"] = "error"
        task["errors"].append(str(e))


# ============== HTML模板 ==============

HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>PDF AI识别工具</title>
    <style>
        * { box-sizing: border-box; margin: 0; padding: 0; }
        body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; background: #f5f5f5; padding: 20px; }
        .container { max-width: 900px; margin: 0 auto; }
        h1 { text-align: center; color: #333; margin-bottom: 30px; }
        
        .card { background: white; border-radius: 12px; padding: 24px; margin-bottom: 20px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); }
        .card h2 { color: #666; font-size: 16px; margin-bottom: 16px; border-bottom: 1px solid #eee; padding-bottom: 10px; }
        
        .form-group { margin-bottom: 16px; }
        .form-group label { display: block; margin-bottom: 6px; color: #555; font-weight: 500; }
        .form-group input[type="text"], .form-group select, .form-group textarea { width: 100%; padding: 10px 12px; border: 1px solid #ddd; border-radius: 6px; font-size: 14px; }
        .form-group textarea { height: 100px; resize: vertical; font-family: inherit; }
        .form-group small { color: #888; margin-top: 4px; display: block; }
        
        .btn { display: inline-block; padding: 12px 24px; border: none; border-radius: 6px; cursor: pointer; font-size: 14px; font-weight: 500; transition: all 0.2s; }
        .btn-primary { background: #4a90d9; color: white; }
        .btn-primary:hover { background: #357abd; }
        .btn-primary:disabled { background: #ccc; cursor: not-allowed; }
        .btn-success { background: #5cb85c; color: white; }
        .btn-success:hover { background: #4cae4c; }
        .btn-danger { background: #d9534f; color: white; }
        .btn-danger:hover { background: #c9302c; }
        .btn-danger:disabled { background: #ccc; cursor: not-allowed; }
        
        .upload-area { border: 2px dashed #ddd; border-radius: 8px; padding: 40px; text-align: center; cursor: pointer; transition: all 0.2s; }
        .upload-area:hover { border-color: #4a90d9; background: #f8f8f8; }
        .upload-area input { display: none; }
        .upload-area .icon { font-size: 48px; margin-bottom: 10px; }
        .upload-area .text { color: #666; }
        .upload-area .text strong { color: #4a90d9; }
        
        .file-list { margin-top: 16px; }
        .file-item { background: #f8f8f8; padding: 12px 16px; border-radius: 6px; margin-bottom: 8px; }
        .file-item .name { font-weight: 500; color: #333; }
        .file-item .info { color: #4a90d9; font-size: 13px; }
        
        .progress-container { margin-top: 20px; display: none; }
        .progress-container.show { display: block; }
        .progress-bar { height: 24px; background: #e0e0e0; border-radius: 12px; overflow: hidden; }
        .progress-bar .fill { height: 100%; background: linear-gradient(90deg, #4a90d9, #5cb85c); transition: width 0.3s; width: 0%; }
        .progress-text { text-align: center; margin-top: 8px; color: #666; font-size: 14px; }
        
        .status { padding: 12px 16px; border-radius: 6px; margin-top: 16px; display: none; }
        .status.show { display: block; }
        .status.running { background: #e3f2fd; color: #1976d2; }
        .status.completed { background: #e8f5e9; color: #388e3c; }
        .status.error { background: #ffebee; color: #c62828; }
        
        .preview-section { margin-top: 20px; display: none; }
        .preview-section.show { display: block; }
        .preview-content { background: #f8f8f8; padding: 16px; border-radius: 6px; max-height: 400px; overflow-y: auto; white-space: pre-wrap; font-size: 13px; line-height: 1.6; }
        
        .actions { display: flex; gap: 12px; margin-top: 16px; }
        
        .model-list { max-height: 200px; overflow-y: auto; border: 1px solid #ddd; border-radius: 6px; padding: 8px; }
        .model-item { padding: 8px 12px; border-radius: 4px; cursor: pointer; transition: background 0.2s; }
        .model-item:hover { background: #f0f0f0; }
        .model-item.selected { background: #e3f2fd; color: #1976d2; }
        
        .preset-select { display: flex; gap: 8px; flex-wrap: wrap; margin-bottom: 12px; }
        .preset-btn { padding: 8px 16px; border: 1px solid #ddd; border-radius: 20px; background: white; cursor: pointer; font-size: 13px; transition: all 0.2s; }
        .preset-btn:hover { border-color: #4a90d9; }
        .preset-btn.active { background: #4a90d9; color: white; border-color: #4a90d9; }
    </style>
</head>
<body>
    <div class="container">
        <h1>📄 PDF AI识别工具</h1>
        
        <!-- LM Studio 配置 -->
        <div class="card">
            <h2>⚙️ LM Studio 配置</h2>
            <div class="form-group">
                <label>API地址</label>
                <input type="text" id="apiUrl" value="http://localhost:1234" placeholder="http://localhost:1234">
            </div>
            <div class="form-group">
                <label>选择模型</label>
                <div style="display: flex; gap: 8px; margin-bottom: 8px;">
                    <button class="btn btn-primary" onclick="loadModels()" style="padding: 8px 16px;">🔄 刷新模型列表</button>
                </div>
                <div class="model-list" id="modelList">
                    <div class="model-item">点击"刷新模型列表"加载可用模型</div>
                </div>
                <input type="hidden" id="selectedModel" value="">
            </div>
        </div>
        
        <!-- 文件上传 -->
        <div class="card">
            <h2>📤 上传PDF文件</h2>
            <div class="upload-area" onclick="document.getElementById('fileInput').click()">
                <input type="file" id="fileInput" accept=".pdf" multiple onchange="handleFileSelect(event)">
                <div class="icon">📄</div>
                <div class="text">点击或拖拽PDF文件到这里<br><strong>支持批量上传多个PDF</strong></div>
            </div>
            <div class="file-list" id="fileList"></div>
        </div>
        
        <!-- 提示词预设 -->
        <div class="card">
            <h2>💬 识别提示词</h2>
            <div class="form-group">
                <label>选择预设模板（点击应用）</label>
                <div class="preset-select" id="presetSelect"></div>
            </div>
            <div class="form-group">
                <label>自定义提示词</label>
                <textarea id="prompt" placeholder="选择一个预设模板，或手动输入提示词..."></textarea>
                <small>提示：工具会原样识别文档内容，不会进行AI总结。</small>
            </div>
        </div>
        
        <!-- 操作按钮 -->
        <div class="card">
            <div class="actions">
                <button class="btn btn-primary" id="startBtn" onclick="startProcessing()" disabled>🚀 开始识别</button>
                <button class="btn btn-danger" id="stopBtn" onclick="stopProcessing()" style="display: none;">🛑 停止</button>
                <button class="btn btn-success" id="downloadDocxBtn" onclick="downloadResult('docx')" style="display: none;">📥 下载Word</button>
                <button class="btn btn-success" id="downloadMdBtn" onclick="downloadResult('md')" style="display: none;">📥 下载MD</button>
            </div>
            
            <div class="progress-container" id="progressContainer">
                <div class="progress-bar"><div class="fill" id="progressFill"></div></div>
                <div class="progress-text" id="progressText">准备中...</div>
            </div>
            
            <div class="status" id="statusBox"></div>
        </div>
        
        <!-- 结果预览 -->
        <div class="card preview-section" id="previewSection">
            <h2>👁️ 结果预览</h2>
            <div style="margin-bottom: 12px;">
                <small>页码: <span id="previewPageNum">1</span></small>
                <button class="btn btn-primary" onclick="prevPage()" style="padding: 6px 12px; margin-left: 8px;">上一页</button>
                <button class="btn btn-primary" onclick="nextPage()" style="padding: 6px 12px;">下一页</button>
            </div>
            <div class="preview-content" id="previewContent">等待处理完成...</div>
        </div>
    </div>
    
    <script>
        let currentTaskId = null;
        let eventSource = null;
        let lastPreviewPage = 0;
        let presets = [];
        
        async function loadModels() {
            const url = document.getElementById('apiUrl').value;
            const modelList = document.getElementById('modelList');
            modelList.innerHTML = '<div class="model-item">加载中...</div>';
            
            try {
                const resp = await fetch(`/api/models?url=${encodeURIComponent(url)}`);
                const data = await resp.json();
                
                if (data.success && data.all_models.length > 0) {
                    const models = data.vision_models || data.all_models;
                    modelList.innerHTML = models.map(m => 
                        `<div class="model-item" onclick="selectModel(this, '${m.replace(/'/g, "\\\\'")}')">${m}</div>`
                    ).join('');
                } else {
                    modelList.innerHTML = '<div class="model-item">未检测到模型，请确保LM Studio正在运行</div>';
                }
            } catch (e) {
                modelList.innerHTML = `<div class="model-item">连接失败: ${e.message}</div>`;
            }
        }
        
        async function loadPresets() {
            try {
                const resp = await fetch('/api/presets');
                const data = await resp.json();
                if (data.success) {
                    presets = data.presets;
                    const container = document.getElementById('presetSelect');
                    container.innerHTML = presets.map(p => 
                        `<div class="preset-btn" onclick="selectPreset('${p.id}')">${p.name}</div>`
                    ).join('');
                }
            } catch (e) {
                console.error('加载预设失败:', e);
            }
        }
        
        function selectPreset(presetId) {
            const preset = presets.find(p => p.id === presetId);
            if (preset) {
                document.getElementById('prompt').value = preset.prompt;
                document.querySelectorAll('.preset-btn').forEach(btn => btn.classList.remove('active'));
                event.target.classList.add('active');
            }
        }
        
        function selectModel(el, model) {
            document.querySelectorAll('.model-item').forEach(item => item.classList.remove('selected'));
            el.classList.add('selected');
            document.getElementById('selectedModel').value = model;
        }
        
        function handleFileSelect(event) {
            const files = Array.from(event.target.files);
            if (!files.length) return;
            
            const fileList = document.getElementById('fileList');
            fileList.innerHTML = '';
            
            let html = '';
            files.forEach((file, index) => {
                html += `<div class="file-item">
                    <div>
                        <div class="name">${file.name}</div>
                        <div class="info">等待上传...</div>
                    </div>
                </div>`;
            });
            fileList.innerHTML = html;
            
            // 上传第一个文件
            const formData = new FormData();
            formData.append('file', files[0]);
            
            fetch('/api/upload', { method: 'POST', body: formData })
                .then(r => r.json())
                .then(data => {
                    if (data.success) {
                        currentTaskId = data.task_id;
                        fileList.querySelector('.info').textContent = `已上传 - 共 ${data.total_pages} 页`;
                        document.getElementById('startBtn').disabled = false;
                    } else {
                        alert('上传失败: ' + data.error);
                    }
                });
        }
        
        function startProcessing() {
            if (!currentTaskId) return;
            
            const model = document.getElementById('selectedModel').value;
            const url = document.getElementById('apiUrl').value;
            const prompt = document.getElementById('prompt').value;
            
            if (!model) {
                alert('请先选择一个模型');
                return;
            }
            
            fetch('/api/start', {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify({
                    task_id: currentTaskId,
                    model: model,
                    url: url,
                    prompt: prompt
                })
            }).then(r => r.json()).then(data => {
                if (data.success) {
                    showProgress();
                } else {
                    alert('启动失败: ' + data.error);
                }
            });
        }
        
        function showProgress() {
            document.getElementById('progressContainer').classList.add('show');
            document.getElementById('startBtn').disabled = true;
            document.getElementById('stopBtn').style.display = 'inline-block';
            document.getElementById('statusBox').className = 'status running show';
            document.getElementById('statusBox').textContent = '处理中...';
            
            if (eventSource) eventSource.close();
            eventSource = new EventSource(`/api/stream/${currentTaskId}`);
            
            eventSource.onmessage = function(e) {
                const data = JSON.parse(e.data);
                updateProgress(data);
            };
            
            eventSource.onerror = function() {
                eventSource.close();
                checkFinalStatus();
            };
        }
        
        function updateProgress(data) {
            document.getElementById('progressFill').style.width = data.progress + '%';
            document.getElementById('progressText').textContent = `第 ${data.current_page} / ${data.total_pages} 页 (${data.progress}%)`;
            
            if (data.done) {
                eventSource.close();
                checkFinalStatus();
            }
        }
        
        function checkFinalStatus() {
            fetch(`/api/status/${currentTaskId}`)
                .then(r => r.json())
                .then(data => {
                    const statusBox = document.getElementById('statusBox');
                    
                    if (data.status === 'completed') {
                        statusBox.className = 'status completed show';
                        statusBox.textContent = '处理完成！';
                        document.getElementById('stopBtn').style.display = 'none';
                        document.getElementById('downloadDocxBtn').style.display = 'inline-block';
                        document.getElementById('downloadMdBtn').style.display = 'inline-block';
                        document.getElementById('previewSection').classList.add('show');
                        loadPreview(0);
                    } else if (data.status === 'error') {
                        statusBox.className = 'status error show';
                        statusBox.textContent = '处理出错: ' + (data.errors || []).join(', ');
                        document.getElementById('stopBtn').style.display = 'none';
                    } else {
                        statusBox.className = 'status running show';
                        statusBox.textContent = `处理中... ${data.progress}%`;
                    }
                });
        }
        
        function loadPreview(page) {
            lastPreviewPage = page;
            document.getElementById('previewPageNum').textContent = page + 1;
            
            fetch(`/api/preview/${currentTaskId}/${page}`)
                .then(r => r.json())
                .then(data => {
                    if (data.success) {
                        document.getElementById('previewContent').textContent = data.content;
                    } else {
                        document.getElementById('previewContent').textContent = '无内容';
                    }
                });
        }
        
        function prevPage() { if (lastPreviewPage > 0) loadPreview(lastPreviewPage - 1); }
        function nextPage() { loadPreview(lastPreviewPage + 1); }
        function downloadResult(fmt) { window.location.href = `/api/download/${currentTaskId}?format=${fmt}`; }
        
        function stopProcessing() {
            if (!currentTaskId) return;
            
            fetch(`/api/stop/${currentTaskId}`, { method: 'POST' })
                .then(r => r.json())
                .then(data => {
                    if (data.success) {
                        document.getElementById('stopBtn').disabled = true;
                        document.getElementById('statusBox').textContent = '已停止';
                        document.getElementById('statusBox').className = 'status error show';
                    }
                });
        }
        
        window.onload = function() {
            loadModels();
            loadPresets();
        };
    </script>
</body>
</html>
"""

if __name__ == "__main__":
    print("=" * 50)
    print("PDF AI识别工具 - 仅支持LM Studio")
    print()
    print("请确保LM Studio正在运行")
    print("默认地址: http://localhost:1234")
    print()
    print("请打开浏览器访问: http://localhost:5000")
    print("=" * 50)
    app.run(host="0.0.0.0", port=5000, debug=False)
